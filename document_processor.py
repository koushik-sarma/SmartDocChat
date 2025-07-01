import os
import logging
from typing import Generator, List, Dict
import pdfplumber
import docx
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Multi-format document processor supporting PDF, DOCX, TXT, and MD files.
    """
    
    def __init__(self, chunk_size: int = 1000):
        self.chunk_size = chunk_size
        self.supported_formats = {'.pdf', '.docx', '.txt', '.md'}
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported."""
        return Path(filename).suffix.lower() in self.supported_formats
    
    def extract_text_chunks(self, file_path: str) -> Generator[str, None, None]:
        """
        Extract text from any supported document format in chunks.
        """
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                yield from self._extract_pdf_text(file_path)
            elif file_ext == '.docx':
                yield from self._extract_docx_text(file_path)
            elif file_ext in ['.txt', '.md']:
                yield from self._extract_text_file(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_pdf_text(self, pdf_path: str) -> Generator[str, None, None]:
        """Extract text from PDF using pdfplumber with PyMuPDF fallback."""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                current_chunk = ""
                
                for page_num, page in enumerate(pdf.pages):
                    try:
                        # Add timeout and error handling for problematic PDFs
                        text = None
                        try:
                            text = page.extract_text()
                        except Exception as extract_error:
                            logger.warning(f"Failed to extract text from page {page_num + 1} with pdfplumber: {extract_error}")
                            continue
                            
                        if text:
                            clean_text = self._clean_text(text)
                            current_chunk += " " + clean_text
                            
                            # Check if chunk is large enough
                            if len(current_chunk.split()) >= self.chunk_size:
                                words = current_chunk.split()
                                chunk = " ".join(words[:self.chunk_size])
                                yield chunk
                                current_chunk = " ".join(words[self.chunk_size:])
                    except Exception as page_error:
                        logger.warning(f"Error on PDF page {page_num + 1}: {page_error}")
                        continue
                
                # Yield remaining chunk
                if current_chunk.strip():
                    yield current_chunk
                    
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            # Fallback to PyMuPDF if available
            try:
                import pymupdf as fitz
                doc = fitz.open(pdf_path)
                current_chunk = ""
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text = page.get_text()
                    if text:
                        clean_text = self._clean_text(text)
                        current_chunk += " " + clean_text
                        
                        if len(current_chunk.split()) >= self.chunk_size:
                            words = current_chunk.split()
                            chunk = " ".join(words[:self.chunk_size])
                            yield chunk
                            current_chunk = " ".join(words[self.chunk_size:])
                
                if current_chunk.strip():
                    yield current_chunk
                    
                doc.close()
            except ImportError:
                raise ValueError(f"Failed to process PDF: {str(e)}")
    
    def _extract_docx_text(self, docx_path: str) -> Generator[str, None, None]:
        """Extract text from DOCX files."""
        try:
            doc = docx.Document(docx_path)
            current_chunk = ""
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    current_chunk += " " + text
                    
                    # Check if chunk is large enough
                    if len(current_chunk.split()) >= self.chunk_size:
                        words = current_chunk.split()
                        chunk = " ".join(words[:self.chunk_size])
                        yield chunk
                        current_chunk = " ".join(words[self.chunk_size:])
            
            # Yield remaining chunk
            if current_chunk.strip():
                yield current_chunk
                
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")
    
    def _extract_text_file(self, file_path: str) -> Generator[str, None, None]:
        """Extract text from plain text and markdown files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                current_chunk = ""
                
                for line in file:
                    line = line.strip()
                    if line:
                        current_chunk += " " + line
                        
                        # Check if chunk is large enough
                        if len(current_chunk.split()) >= self.chunk_size:
                            words = current_chunk.split()
                            chunk = " ".join(words[:self.chunk_size])
                            yield chunk
                            current_chunk = " ".join(words[self.chunk_size:])
                
                # Yield remaining chunk
                if current_chunk.strip():
                    yield current_chunk
                    
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                    clean_content = self._clean_text(content)
                    words = clean_content.split()
                    
                    for i in range(0, len(words), self.chunk_size):
                        chunk = " ".join(words[i:i + self.chunk_size])
                        yield chunk
            except Exception as e:
                logger.error(f"Text file extraction failed: {e}")
                raise ValueError(f"Failed to process text file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text while preserving chemical equations and special characters."""
        import re
        
        # Convert Unicode special characters first
        text = self._convert_special_characters(text)
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Only remove non-printable control characters, keep special chars
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        return text.strip()
    
    def _convert_special_characters(self, text: str) -> str:
        """Convert Unicode superscripts, subscripts, and special characters to readable format."""
        # Chemical formula mappings
        subscript_map = {
            '₀': '0', '₁': '1', '₂': '2', '₃': '3', '₄': '4', '₅': '5', '₆': '6', '₇': '7', '₈': '8', '₉': '9',
            'ₐ': 'a', 'ₑ': 'e', 'ᵢ': 'i', 'ₒ': 'o', 'ᵤ': 'u', 'ₓ': 'x', 'ₙ': 'n', 'ₘ': 'm', 'ₚ': 'p', 'ₛ': 's', 'ₜ': 't'
        }
        
        superscript_map = {
            '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4', '⁵': '5', '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9',
            'ᵃ': 'a', 'ᵇ': 'b', 'ᶜ': 'c', 'ᵈ': 'd', 'ᵉ': 'e', 'ᶠ': 'f', 'ᵍ': 'g', 'ʰ': 'h', 'ⁱ': 'i', 'ʲ': 'j',
            'ᵏ': 'k', 'ˡ': 'l', 'ᵐ': 'm', 'ⁿ': 'n', 'ᵒ': 'o', 'ᵖ': 'p', 'ʳ': 'r', 'ˢ': 's', 'ᵗ': 't', 'ᵘ': 'u',
            'ᵛ': 'v', 'ʷ': 'w', 'ˣ': 'x', 'ʸ': 'y', 'ᶻ': 'z', '⁺': '+', '⁻': '-', '⁼': '='
        }
        
        # Apply subscript conversions with underscore notation
        for sub, normal in subscript_map.items():
            text = text.replace(sub, f'_{normal}')
        
        # Apply superscript conversions with caret notation
        for sup, normal in superscript_map.items():
            text = text.replace(sup, f'^{normal}')
        
        return text
    
    def get_document_info(self, file_path: str) -> dict:
        """Get basic information about the document."""
        file_ext = Path(file_path).suffix.lower()
        
        try:
            info = {
                'filename': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'format': file_ext,
                'pages': 0
            }
            
            if file_ext == '.pdf':
                with pdfplumber.open(file_path) as pdf:
                    info['pages'] = len(pdf.pages)
            elif file_ext == '.docx':
                doc = docx.Document(file_path)
                info['pages'] = len(doc.paragraphs)
            elif file_ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as file:
                    info['pages'] = len(file.readlines())
            
            return info
            
        except Exception as e:
            logger.error(f"Error getting document info: {str(e)}")
            raise ValueError(f"Cannot read document file: {str(e)}")
    
    def compare_documents(self, doc_paths: List[str]) -> Dict:
        """
        Compare multiple documents and find common themes/differences.
        """
        if len(doc_paths) < 2:
            return {'error': 'Need at least 2 documents to compare'}
        
        try:
            doc_contents = {}
            
            # Extract text from each document
            for doc_path in doc_paths:
                filename = os.path.basename(doc_path)
                chunks = list(self.extract_text_chunks(doc_path))
                doc_contents[filename] = ' '.join(chunks)
            
            # Basic comparison metrics
            comparison = {
                'documents': list(doc_contents.keys()),
                'word_counts': {},
                'common_themes': [],
                'unique_content': {}
            }
            
            # Calculate word counts
            for filename, content in doc_contents.items():
                words = content.split()
                comparison['word_counts'][filename] = len(words)
            
            # Find common and unique words (simplified approach)
            all_words = set()
            doc_words = {}
            
            for filename, content in doc_contents.items():
                words = set(word.lower() for word in content.split() if len(word) > 3)
                doc_words[filename] = words
                all_words.update(words)
            
            # Find common words across all documents
            common_words = set.intersection(*doc_words.values()) if doc_words else set()
            comparison['common_themes'] = list(common_words)[:20]  # Top 20 common themes
            
            # Find unique content per document
            for filename, words in doc_words.items():
                other_words = set()
                for other_file, other_word_set in doc_words.items():
                    if other_file != filename:
                        other_words.update(other_word_set)
                
                unique_words = words - other_words
                comparison['unique_content'][filename] = list(unique_words)[:15]  # Top 15 unique words
            
            return comparison
            
        except Exception as e:
            logger.error(f"Error comparing documents: {e}")
            return {'error': f'Document comparison failed: {str(e)}'}