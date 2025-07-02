"""
Unified document processing service.
Handles all document-related operations including upload, processing, and management.
"""

import os
import logging
import tempfile
from typing import List, Dict, Generator, Optional, Tuple
from datetime import datetime
from werkzeug.utils import secure_filename
from werkzeug.datastructures import FileStorage

from app import db
from models import Document
from .base_service import BaseService

# Import processors
try:
    import pdfplumber
    from docx import Document as DocxDocument
except ImportError as e:
    logging.warning(f"Optional dependency missing: {e}")

logger = logging.getLogger(__name__)

class DocumentService(BaseService):
    """Unified service for all document operations."""
    
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'md'}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    CHUNK_SIZE = 1000  # Words per chunk
    
    def __init__(self):
        super().__init__()
        self.upload_folder = 'uploads'
        os.makedirs(self.upload_folder, exist_ok=True)
    
    def is_allowed_file(self, filename: str) -> bool:
        """Check if file type is supported."""
        return ('.' in filename and 
                filename.rsplit('.', 1)[1].lower() in self.ALLOWED_EXTENSIONS)
    
    def upload_document(self, file: FileStorage, session_id: str) -> Dict:
        """
        Upload and process a document.
        Returns: {success: bool, document_id: int, message: str, chunk_count: int}
        """
        try:
            # Validate file
            if not file or not file.filename:
                return self.error_response("No file provided")
            
            if not self.is_allowed_file(file.filename):
                return self.error_response(f"File type not supported. Allowed: {', '.join(self.ALLOWED_EXTENSIONS)}")
            
            # Check file size
            file.seek(0, 2)  # Seek to end
            file_size = file.tell()
            file.seek(0)  # Reset to beginning
            
            if file_size > self.MAX_FILE_SIZE:
                return self.error_response(f"File too large. Maximum size: {self.MAX_FILE_SIZE // (1024*1024)}MB")
            
            # Generate secure filename
            timestamp = int(datetime.now().timestamp())
            secure_name = secure_filename(file.filename)
            filename = f"{timestamp}_{secure_name}"
            file_path = os.path.join(self.upload_folder, filename)
            
            # Save file
            file.save(file_path)
            
            # Process document and count chunks
            chunk_count = 0
            try:
                for _ in self.extract_text_chunks(file_path):
                    chunk_count += 1
                
                # Save to database
                document = Document(
                    session_id=session_id,
                    filename=secure_name,
                    file_path=file_path,
                    chunk_count=chunk_count,
                    file_size=file_size,
                    is_active=True
                )
                
                db.session.add(document)
                db.session.commit()
                
                return self.success_response({
                    'document_id': document.id,
                    'filename': secure_name,
                    'chunk_count': chunk_count,
                    'file_size': file_size,
                    'message': f"Successfully uploaded {secure_name}. Note: Vector embeddings temporarily disabled due to API quota."
                })
                
            except Exception as e:
                # Clean up file if processing failed
                if os.path.exists(file_path):
                    os.remove(file_path)
                logger.error(f"Document processing failed: {e}")
                return self.error_response(f"Failed to process document: {str(e)}")
                
        except Exception as e:
            logger.error(f"Document upload failed: {e}")
            return self.error_response(f"Upload failed: {str(e)}")
    
    def extract_text_chunks(self, file_path: str) -> Generator[str, None, None]:
        """Extract text from any supported document format in chunks."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.lower().split('.')[-1]
        
        try:
            if file_ext == 'pdf':
                yield from self._extract_pdf_chunks(file_path)
            elif file_ext == 'docx':
                yield from self._extract_docx_chunks(file_path)
            elif file_ext in ['txt', 'md']:
                yield from self._extract_text_chunks(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            raise
    
    def _extract_pdf_chunks(self, pdf_path: str) -> Generator[str, None, None]:
        """Extract text from PDF with optimized processing for large files."""
        file_size = os.path.getsize(pdf_path)
        logger.info(f"Processing PDF: {os.path.basename(pdf_path)} ({file_size} bytes)")
        
        # Process with optimized settings for large files
        try:
            with pdfplumber.open(pdf_path) as pdf:
                current_chunk = ""
                word_count = 0
                total_pages = len(pdf.pages)
                logger.info(f"Processing {total_pages} pages from PDF")
                
                for page_num, page in enumerate(pdf.pages):
                    try:
                        # Log progress for large files
                        if page_num % 50 == 0:
                            logger.info(f"Processing page {page_num + 1}/{total_pages}")
                        
                        text = page.extract_text() or ""
                        cleaned_text = self._clean_text(text)
                        
                        words = cleaned_text.split()
                        for word in words:
                            current_chunk += " " + word
                            word_count += 1
                            
                            if word_count >= self.CHUNK_SIZE:
                                yield current_chunk.strip()
                                current_chunk = ""
                                word_count = 0
                    except Exception as e:
                        logger.warning(f"Error processing page {page_num} in {pdf_path}: {e}")
                        continue
                
                if current_chunk.strip():
                    yield current_chunk.strip()
                return
                    
        except Exception as e:
            logger.error(f"PDF processing failed for {pdf_path}: {e}")
            raise ValueError(f"Could not extract text from PDF: {str(e)}")
    
    def _extract_docx_chunks(self, docx_path: str) -> Generator[str, None, None]:
        """Extract text from DOCX files."""
        try:
            doc = DocxDocument(docx_path)
            current_chunk = ""
            word_count = 0
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if text.strip():
                    cleaned_text = self._clean_text(text)
                    words = cleaned_text.split()
                    
                    for word in words:
                        current_chunk += " " + word
                        word_count += 1
                        
                        if word_count >= self.CHUNK_SIZE:
                            yield current_chunk.strip()
                            current_chunk = ""
                            word_count = 0
            
            if current_chunk.strip():
                yield current_chunk.strip()
                
        except Exception as e:
            logger.error(f"DOCX extraction failed for {docx_path}: {e}")
            raise ValueError(f"Could not extract text from DOCX: {str(e)}")
    
    def _extract_text_chunks(self, file_path: str) -> Generator[str, None, None]:
        """Extract text from plain text and markdown files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                current_chunk = ""
                word_count = 0
                
                for line in file:
                    cleaned_line = self._clean_text(line)
                    words = cleaned_line.split()
                    
                    for word in words:
                        current_chunk += " " + word
                        word_count += 1
                        
                        if word_count >= self.CHUNK_SIZE:
                            yield current_chunk.strip()
                            current_chunk = ""
                            word_count = 0
                
                if current_chunk.strip():
                    yield current_chunk.strip()
                    
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                    cleaned_content = self._clean_text(content)
                    words = cleaned_content.split()
                    
                    for i in range(0, len(words), self.CHUNK_SIZE):
                        chunk = " ".join(words[i:i + self.CHUNK_SIZE])
                        yield chunk
            except Exception as e:
                logger.error(f"Text file extraction failed: {e}")
                raise ValueError(f"Failed to process text file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text while preserving important formatting."""
        import re
        
        # Convert Unicode special characters
        text = self._convert_special_characters(text)
        
        # Remove excessive whitespace but preserve structure
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters but keep printable text
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        return text.strip()
    
    def _convert_special_characters(self, text: str) -> str:
        """Convert Unicode special characters to readable format."""
        # Chemical equation conversions
        conversions = {
            # Superscripts (common in chemistry)
            '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4', '⁵': '5',
            '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9', '⁺': '+', '⁻': '-',
            
            # Subscripts (common in chemical formulas)
            '₀': '0', '₁': '1', '₂': '2', '₃': '3', '₄': '4', '₅': '5',
            '₆': '6', '₇': '7', '₈': '8', '₉': '9', '₊': '+', '₋': '-',
            
            # Greek letters (common in equations)
            'α': 'alpha', 'β': 'beta', 'γ': 'gamma', 'δ': 'delta',
            'ε': 'epsilon', 'θ': 'theta', 'λ': 'lambda', 'μ': 'mu',
            'π': 'pi', 'ρ': 'rho', 'σ': 'sigma', 'φ': 'phi', 'ω': 'omega',
            
            # Mathematical symbols
            '×': 'x', '÷': '/', '≈': '~', '≡': '=', '≠': '!=',
            '≤': '<=', '≥': '>=', '∞': 'infinity',
        }
        
        for unicode_char, replacement in conversions.items():
            text = text.replace(unicode_char, replacement)
        
        return text
    
    def get_session_documents(self, session_id: str) -> List[Document]:
        """Get all documents for a session."""
        return Document.query.filter_by(session_id=session_id).all()
    
    def get_active_documents(self, session_id: str) -> List[Document]:
        """Get all active documents for a session."""
        return Document.query.filter_by(session_id=session_id, is_active=True).all()
    
    def toggle_document_status(self, document_id: int, session_id: str) -> Dict:
        """Toggle document active/inactive status."""
        try:
            document = Document.query.filter_by(id=document_id, session_id=session_id).first()
            if not document:
                return self.error_response("Document not found")
            
            document.is_active = not document.is_active
            db.session.commit()
            
            status = "enabled" if document.is_active else "disabled"
            return self.success_response({
                'document_id': document_id,
                'is_active': document.is_active,
                'message': f"Document {status} successfully"
            })
            
        except Exception as e:
            logger.error(f"Error toggling document status: {e}")
            return self.error_response("Failed to update document status")
    
    def delete_document(self, document_id: int, session_id: str) -> Dict:
        """Delete a document and its file."""
        try:
            document = Document.query.filter_by(id=document_id, session_id=session_id).first()
            if not document:
                return self.error_response("Document not found")
            
            # Delete file if it exists
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
            
            # Delete from database
            db.session.delete(document)
            db.session.commit()
            
            return self.success_response({
                'document_id': document_id,
                'message': f"Document {document.filename} deleted successfully"
            })
            
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            return self.error_response("Failed to delete document")
    
    def get_document_stats(self, session_id: str) -> Dict:
        """Get statistics about documents for a session."""
        try:
            documents = self.get_session_documents(session_id)
            active_docs = [d for d in documents if d.is_active]
            
            total_chunks = sum(doc.chunk_count for doc in active_docs)
            total_size = sum(doc.file_size for doc in documents)
            
            return {
                'total_documents': len(documents),
                'active_documents': len(active_docs),
                'total_chunks': total_chunks,
                'total_size_mb': round(total_size / (1024 * 1024), 2)
            }
            
        except Exception as e:
            logger.error(f"Error getting document stats: {e}")
            return {
                'total_documents': 0,
                'active_documents': 0,
                'total_chunks': 0,
                'total_size_mb': 0
            }